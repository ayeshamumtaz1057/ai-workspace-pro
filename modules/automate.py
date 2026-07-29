"""Automation & Export — one-click document workflows and session exports.

Ties the other tools together: turn any AI output into a downloadable report,
batch-summarize several PDFs at once, and export your activity history.
"""
import io
from datetime import datetime
import pandas as pd
import streamlit as st
from core import ai, db


def _md_to_docx_bytes(title: str, body: str) -> bytes | None:
    try:
        import docx
        doc = docx.Document()
        doc.add_heading(title, level=0)
        for line in body.splitlines():
            s = line.strip()
            if not s:
                continue
            if s.startswith("### "):
                doc.add_heading(s[4:], level=3)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=2)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif s.startswith(("- ", "* ")):
                doc.add_paragraph(s[2:], style="List Bullet")
            else:
                doc.add_paragraph(s)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return None


def render():
    st.subheader("📤 Automation Center")
    st.caption("Turn AI output into documents, batch-process files and export your data.")

    gen, batch, history = st.tabs(
        ["Report Generator", "Batch PDF Summaries", "Export Activity"])

    # ---- 1. Report Generator --------------------------------------------
    with gen:
        st.markdown("**Generate a polished report from any prompt.**")
        topic = st.text_input("Report topic or instruction",
                              "Executive summary of Q1 sales performance")
        fmt = st.radio("Structure", ["Executive summary", "Detailed report",
                                     "Bullet brief"], horizontal=True)
        if st.button("Generate report", type="primary") and topic.strip():
            with st.spinner("Writing…"):
                body = ai.ask(
                    topic,
                    system=(f"Write a professional {fmt.lower()} in clean markdown with "
                            "headings and bullet points where appropriate."))
            st.markdown(body)
            title = topic[:60]
            c1, c2 = st.columns(2)
            c1.download_button("Download Markdown", body,
                               file_name="report.md", use_container_width=True)
            docx_bytes = _md_to_docx_bytes(title, body)
            if docx_bytes:
                c2.download_button("Download Word (.docx)", docx_bytes,
                                   file_name="report.docx", use_container_width=True)
            db.bump("reports_exported")
            db.log_message("Automation & Export", title, "user", topic)

    # ---- 2. Batch PDF Summaries -----------------------------------------
    with batch:
        st.markdown("**Summarize several PDFs at once into one table.**")
        files = st.file_uploader("Upload PDFs", type=["pdf"],
                                 accept_multiple_files=True, key="batch_pdfs")
        if files and st.button("Summarize all", type="primary"):
            from pypdf import PdfReader
            rows = []
            bar = st.progress(0.0)
            for i, f in enumerate(files, 1):
                try:
                    text = "\n".join((p.extract_text() or "") for p in PdfReader(f).pages)
                    summary = ai.ask(text[:8000],
                                     system="Summarize this document in exactly 2 sentences.")
                except Exception as exc:
                    summary = f"Could not read: {exc}"
                rows.append({"File": f.name, "Summary": summary})
                bar.progress(i / len(files))
                db.bump("pdfs_processed")
            out = pd.DataFrame(rows)
            st.dataframe(out, use_container_width=True)
            st.download_button("Download CSV", out.to_csv(index=False).encode(),
                               "pdf_summaries.csv", "text/csv")

    # ---- 3. Export Activity ---------------------------------------------
    with history:
        st.markdown("**Export your usage history and stats.**")
        stats = db.get_stats()
        st.dataframe(pd.DataFrame([stats]), use_container_width=True)

        recents = db.recent_chats(50)
        if recents:
            df = pd.DataFrame(recents)
            st.download_button("Download activity (CSV)",
                               df.to_csv(index=False).encode(),
                               "activity.csv", "text/csv")

            buf = io.BytesIO()
            with pd.ExcelWriter(buf, engine="openpyxl") as xl:
                pd.DataFrame([stats]).to_excel(xl, index=False, sheet_name="Stats")
                df.to_excel(xl, index=False, sheet_name="Activity")
            st.download_button("Download full report (Excel)", buf.getvalue(),
                               f"workspace_report_{datetime.now():%Y%m%d}.xlsx")
        else:
            st.info("No activity recorded yet. Use the other tools to build history.")
